# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\User\Desktop\School\Python\lab3and4\Designs\frontend.ui'
#
# Created by: PyQt5 UI code generator 5.15.4
#
# WARNING: Any manual changes made to this file will be lost when pyuic5 is
# run again.  Do not edit this file unless you know what you are doing.


from PyQt5 import QtCore, QtGui, QtWidgets
import controller
from tkinter import messagebox
from docx import Document


import pandas
from datetime import date
from docx.shared import Inches



class Ui_Dialog(object):
    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(966, 532)
        self.tabWidget = QtWidgets.QTabWidget(Dialog)
        self.tabWidget.setGeometry(QtCore.QRect(0, 10, 961, 521))
        self.tabWidget.setObjectName("tabWidget")
        self.tab = QtWidgets.QWidget()
        self.tab.setObjectName("tab")
        self.label = QtWidgets.QLabel(self.tab)
        self.label.setGeometry(QtCore.QRect(90, 20, 61, 16))
        self.label.setObjectName("label")
        self.addDoctorName = QtWidgets.QLineEdit(self.tab)
        self.addDoctorName.setGeometry(QtCore.QRect(90, 50, 113, 22))
        self.addDoctorName.setObjectName("addDoctorName")
        self.addDoctorSpecialty = QtWidgets.QLineEdit(self.tab)
        self.addDoctorSpecialty.setGeometry(QtCore.QRect(90, 90, 113, 22))
        self.addDoctorSpecialty.setObjectName("addDoctorSpecialty")
        self.addCost = QtWidgets.QLineEdit(self.tab)
        self.addCost.setGeometry(QtCore.QRect(90, 130, 113, 22))
        self.addCost.setObjectName("addCost")
        self.addNameOfReception = QtWidgets.QLineEdit(self.tab)
        self.addNameOfReception.setGeometry(QtCore.QRect(90, 170, 113, 22))
        self.addNameOfReception.setObjectName("addNameOfReception")
        self.label_2 = QtWidgets.QLabel(self.tab)
        self.label_2.setGeometry(QtCore.QRect(20, 50, 61, 20))
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.tab)
        self.label_3.setGeometry(QtCore.QRect(20, 90, 51, 20))
        self.label_3.setObjectName("label_3")
        self.label_4 = QtWidgets.QLabel(self.tab)
        self.label_4.setGeometry(QtCore.QRect(20, 130, 55, 16))
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.tab)
        self.label_5.setGeometry(QtCore.QRect(20, 170, 55, 16))
        self.label_5.setObjectName("label_5")
        self.addUserButton = QtWidgets.QPushButton(self.tab)
        self.addUserButton.setGeometry(QtCore.QRect(30, 290, 171, 41))
        self.addUserButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.addUserButton.setObjectName("addUserButton")
        self.updateUserButton = QtWidgets.QPushButton(self.tab)
        self.updateUserButton.setGeometry(QtCore.QRect(240, 290, 191, 41))
        self.updateUserButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.updateUserButton.setObjectName("updateUserButton")
        self.label_6 = QtWidgets.QLabel(self.tab)
        self.label_6.setGeometry(QtCore.QRect(300, 20, 91, 16))
        self.label_6.setObjectName("label_6")
        self.deleteUserButton = QtWidgets.QPushButton(self.tab)
        self.deleteUserButton.setGeometry(QtCore.QRect(470, 290, 181, 41))
        self.deleteUserButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.deleteUserButton.setObjectName("deleteUserButton")
        self.tableDoctor = QtWidgets.QTableWidget(self.tab)
        self.tableDoctor.setGeometry(QtCore.QRect(480, 0, 471, 281))
        self.tableDoctor.setSelectionMode(QtWidgets.QAbstractItemView.NoSelection)
        self.tableDoctor.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectColumns)
        self.tableDoctor.setObjectName("tableDoctor")
        self.tableDoctor.setColumnCount(0)
        self.tableDoctor.setRowCount(0)
        self.addPercentage = QtWidgets.QLineEdit(self.tab)
        self.addPercentage.setGeometry(QtCore.QRect(90, 210, 113, 22))
        self.addPercentage.setObjectName("addPercentage")
        self.label_17 = QtWidgets.QLabel(self.tab)
        self.label_17.setGeometry(QtCore.QRect(20, 210, 55, 16))
        self.label_17.setObjectName("label_17")
        self.updatePercentage = QtWidgets.QLineEdit(self.tab)
        self.updatePercentage.setGeometry(QtCore.QRect(310, 243, 113, 22))
        self.updatePercentage.setObjectName("updatePercentage")
        self.updateNameOfReception = QtWidgets.QLineEdit(self.tab)
        self.updateNameOfReception.setGeometry(QtCore.QRect(310, 203, 113, 22))
        self.updateNameOfReception.setObjectName("updateNameOfReception")
        self.label_8 = QtWidgets.QLabel(self.tab)
        self.label_8.setGeometry(QtCore.QRect(240, 203, 55, 16))
        self.label_8.setObjectName("label_8")
        self.label_18 = QtWidgets.QLabel(self.tab)
        self.label_18.setGeometry(QtCore.QRect(240, 243, 55, 16))
        self.label_18.setObjectName("label_18")
        self.updateDoctorName = QtWidgets.QLineEdit(self.tab)
        self.updateDoctorName.setGeometry(QtCore.QRect(310, 83, 113, 22))
        self.updateDoctorName.setObjectName("updateDoctorName")
        self.label_9 = QtWidgets.QLabel(self.tab)
        self.label_9.setGeometry(QtCore.QRect(240, 163, 55, 16))
        self.label_9.setObjectName("label_9")
        self.label_10 = QtWidgets.QLabel(self.tab)
        self.label_10.setGeometry(QtCore.QRect(240, 83, 61, 20))
        self.label_10.setObjectName("label_10")
        self.updateDoctorSpecialty = QtWidgets.QLineEdit(self.tab)
        self.updateDoctorSpecialty.setGeometry(QtCore.QRect(310, 123, 113, 22))
        self.updateDoctorSpecialty.setObjectName("updateDoctorSpecialty")
        self.updateCost = QtWidgets.QLineEdit(self.tab)
        self.updateCost.setGeometry(QtCore.QRect(310, 163, 113, 22))
        self.updateCost.setObjectName("updateCost")
        self.label_11 = QtWidgets.QLabel(self.tab)
        self.label_11.setGeometry(QtCore.QRect(240, 123, 51, 20))
        self.label_11.setObjectName("label_11")
        self.IDDoctor = QtWidgets.QLineEdit(self.tab)
        self.IDDoctor.setGeometry(QtCore.QRect(310, 50, 113, 22))
        self.IDDoctor.setObjectName("IDDoctor")
        self.label_12 = QtWidgets.QLabel(self.tab)
        self.label_12.setGeometry(QtCore.QRect(240, 50, 61, 20))
        self.label_12.setObjectName("label_12")
        self.tabWidget.addTab(self.tab, "")
        self.tab_2 = QtWidgets.QWidget()
        self.tab_2.setObjectName("tab_2")
        self.label_50 = QtWidgets.QLabel(self.tab_2)
        self.label_50.setGeometry(QtCore.QRect(30, 60, 51, 20))
        self.label_50.setObjectName("label_50")
        self.label_52 = QtWidgets.QLabel(self.tab_2)
        self.label_52.setGeometry(QtCore.QRect(30, 197, 55, 16))
        self.label_52.setObjectName("label_52")
        self.deletePatientButton = QtWidgets.QPushButton(self.tab_2)
        self.deletePatientButton.setGeometry(QtCore.QRect(470, 300, 181, 41))
        self.deletePatientButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.deletePatientButton.setObjectName("deletePatientButton")
        self.label_51 = QtWidgets.QLabel(self.tab_2)
        self.label_51.setGeometry(QtCore.QRect(30, 100, 71, 20))
        self.label_51.setObjectName("label_51")
        self.addPatientsName = QtWidgets.QLineEdit(self.tab_2)
        self.addPatientsName.setGeometry(QtCore.QRect(90, 60, 113, 22))
        self.addPatientsName.setObjectName("addPatientsName")
        self.addPatientsSurname = QtWidgets.QLineEdit(self.tab_2)
        self.addPatientsSurname.setGeometry(QtCore.QRect(90, 100, 113, 22))
        self.addPatientsSurname.setObjectName("addPatientsSurname")
        self.updatePatientButton = QtWidgets.QPushButton(self.tab_2)
        self.updatePatientButton.setGeometry(QtCore.QRect(240, 300, 191, 41))
        self.updatePatientButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.updatePatientButton.setObjectName("updatePatientButton")
        self.addPatientsAdress = QtWidgets.QLineEdit(self.tab_2)
        self.addPatientsAdress.setGeometry(QtCore.QRect(30, 220, 171, 31))
        self.addPatientsAdress.setObjectName("addPatientsAdress")
        self.label_49 = QtWidgets.QLabel(self.tab_2)
        self.label_49.setGeometry(QtCore.QRect(90, 30, 61, 16))
        self.label_49.setObjectName("label_49")
        self.addPatientButton = QtWidgets.QPushButton(self.tab_2)
        self.addPatientButton.setGeometry(QtCore.QRect(20, 300, 181, 41))
        self.addPatientButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.addPatientButton.setObjectName("addPatientButton")
        self.tablePatients = QtWidgets.QTableWidget(self.tab_2)
        self.tablePatients.setGeometry(QtCore.QRect(490, 10, 461, 281))
        self.tablePatients.setSelectionMode(QtWidgets.QAbstractItemView.NoSelection)
        self.tablePatients.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectColumns)
        self.tablePatients.setObjectName("tablePatients")
        self.tablePatients.setColumnCount(0)
        self.tablePatients.setRowCount(0)
        self.addPatientsBirthday = QtWidgets.QDateEdit(self.tab_2)
        self.addPatientsBirthday.setGeometry(QtCore.QRect(90, 150, 110, 22))
        self.addPatientsBirthday.setCalendarPopup(True)
        self.addPatientsBirthday.setTimeSpec(QtCore.Qt.LocalTime)
        self.addPatientsBirthday.setObjectName("addPatientsBirthday")
        self.label_53 = QtWidgets.QLabel(self.tab_2)
        self.label_53.setGeometry(QtCore.QRect(30, 150, 55, 16))
        self.label_53.setObjectName("label_53")
        self.label_54 = QtWidgets.QLabel(self.tab_2)
        self.label_54.setGeometry(QtCore.QRect(250, 95, 51, 20))
        self.label_54.setObjectName("label_54")
        self.updatePatientsSurname = QtWidgets.QLineEdit(self.tab_2)
        self.updatePatientsSurname.setGeometry(QtCore.QRect(310, 135, 113, 22))
        self.updatePatientsSurname.setObjectName("updatePatientsSurname")
        self.label_56 = QtWidgets.QLabel(self.tab_2)
        self.label_56.setGeometry(QtCore.QRect(310, 30, 91, 16))
        self.label_56.setObjectName("label_56")
        self.updatePatientsBirthday = QtWidgets.QDateEdit(self.tab_2)
        self.updatePatientsBirthday.setGeometry(QtCore.QRect(310, 185, 110, 22))
        self.updatePatientsBirthday.setCalendarPopup(True)
        self.updatePatientsBirthday.setTimeSpec(QtCore.Qt.LocalTime)
        self.updatePatientsBirthday.setObjectName("updatePatientsBirthday")
        self.updatePatientsAdress = QtWidgets.QLineEdit(self.tab_2)
        self.updatePatientsAdress.setGeometry(QtCore.QRect(250, 255, 171, 31))
        self.updatePatientsAdress.setObjectName("updatePatientsAdress")
        self.updatePatientsName = QtWidgets.QLineEdit(self.tab_2)
        self.updatePatientsName.setGeometry(QtCore.QRect(310, 95, 113, 22))
        self.updatePatientsName.setObjectName("updatePatientsName")
        self.label_57 = QtWidgets.QLabel(self.tab_2)
        self.label_57.setGeometry(QtCore.QRect(250, 232, 55, 16))
        self.label_57.setObjectName("label_57")
        self.label_58 = QtWidgets.QLabel(self.tab_2)
        self.label_58.setGeometry(QtCore.QRect(250, 185, 55, 16))
        self.label_58.setObjectName("label_58")
        self.label_59 = QtWidgets.QLabel(self.tab_2)
        self.label_59.setGeometry(QtCore.QRect(250, 135, 71, 20))
        self.label_59.setObjectName("label_59")
        self.IDPatient = QtWidgets.QLineEdit(self.tab_2)
        self.IDPatient.setGeometry(QtCore.QRect(310, 60, 113, 22))
        self.IDPatient.setObjectName("IDPatient")
        self.label_55 = QtWidgets.QLabel(self.tab_2)
        self.label_55.setGeometry(QtCore.QRect(250, 60, 51, 20))
        self.label_55.setObjectName("label_55")
        self.tabWidget.addTab(self.tab_2, "")
        self.tab_3 = QtWidgets.QWidget()
        self.tab_3.setObjectName("tab_3")
        self.histiroyIdDoctor = QtWidgets.QLabel(self.tab_3)
        self.histiroyIdDoctor.setGeometry(QtCore.QRect(20, 50, 21, 16))
        self.histiroyIdDoctor.setObjectName("histiroyIdDoctor")
        self.historyNamePatient = QtWidgets.QLabel(self.tab_3)
        self.historyNamePatient.setGeometry(QtCore.QRect(190, 90, 91, 16))
        self.historyNamePatient.setObjectName("historyNamePatient")
        self.historySpecialtyDoctor = QtWidgets.QLabel(self.tab_3)
        self.historySpecialtyDoctor.setGeometry(QtCore.QRect(20, 130, 55, 16))
        self.historySpecialtyDoctor.setObjectName("historySpecialtyDoctor")
        self.histiroyNameDoctor = QtWidgets.QLabel(self.tab_3)
        self.histiroyNameDoctor.setGeometry(QtCore.QRect(20, 90, 71, 20))
        self.histiroyNameDoctor.setObjectName("histiroyNameDoctor")
        self.historyCost = QtWidgets.QLabel(self.tab_3)
        self.historyCost.setGeometry(QtCore.QRect(20, 170, 41, 16))
        self.historyCost.setObjectName("historyCost")
        self.historyPatientBirthday = QtWidgets.QLabel(self.tab_3)
        self.historyPatientBirthday.setGeometry(QtCore.QRect(190, 170, 71, 16))
        self.historyPatientBirthday.setObjectName("historyPatientBirthday")
        self.historyCost_2 = QtWidgets.QLabel(self.tab_3)
        self.historyCost_2.setGeometry(QtCore.QRect(20, 250, 51, 16))
        self.historyCost_2.setObjectName("historyCost_2")
        self.historyPatientAdress = QtWidgets.QLabel(self.tab_3)
        self.historyPatientAdress.setGeometry(QtCore.QRect(190, 200, 161, 41))
        self.historyPatientAdress.setObjectName("historyPatientAdress")
        self.historyPatientSurname = QtWidgets.QLabel(self.tab_3)
        self.historyPatientSurname.setGeometry(QtCore.QRect(190, 130, 101, 16))
        self.historyPatientSurname.setObjectName("historyPatientSurname")
        self.selectDatas = QtWidgets.QPushButton(self.tab_3)
        self.selectDatas.setGeometry(QtCore.QRect(10, 430, 171, 41))
        self.selectDatas.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.selectDatas.setObjectName("selectDatas")
        self.label_78 = QtWidgets.QLabel(self.tab_3)
        self.label_78.setGeometry(QtCore.QRect(20, 20, 71, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_78.setFont(font)
        self.label_78.setObjectName("label_78")
        self.historyReception = QtWidgets.QLabel(self.tab_3)
        self.historyReception.setGeometry(QtCore.QRect(20, 210, 51, 20))
        self.historyReception.setObjectName("historyReception")
        self.tableWidget_3 = QtWidgets.QTableWidget(self.tab_3)
        self.tableWidget_3.setGeometry(QtCore.QRect(400, 280, 531, 191))
        self.tableWidget_3.setSelectionMode(QtWidgets.QAbstractItemView.NoSelection)
        self.tableWidget_3.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectColumns)
        self.tableWidget_3.setObjectName("tableWidget_3")
        self.tableWidget_3.setColumnCount(0)
        self.tableWidget_3.setRowCount(0)
        self.label_79 = QtWidgets.QLabel(self.tab_3)
        self.label_79.setGeometry(QtCore.QRect(190, 20, 71, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_79.setFont(font)
        self.label_79.setObjectName("label_79")
        self.historyIdPatient = QtWidgets.QLabel(self.tab_3)
        self.historyIdPatient.setGeometry(QtCore.QRect(190, 50, 21, 16))
        self.historyIdPatient.setObjectName("historyIdPatient")
        self.saveData = QtWidgets.QPushButton(self.tab_3)
        self.saveData.setGeometry(QtCore.QRect(210, 430, 171, 41))
        self.saveData.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.saveData.setObjectName("saveData")
        self.tableWidget_4 = QtWidgets.QTableWidget(self.tab_3)
        self.tableWidget_4.setGeometry(QtCore.QRect(400, 30, 531, 211))
        self.tableWidget_4.setSelectionMode(QtWidgets.QAbstractItemView.NoSelection)
        self.tableWidget_4.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectColumns)
        self.tableWidget_4.setObjectName("tableWidget_4")
        self.tableWidget_4.setColumnCount(0)
        self.tableWidget_4.setRowCount(0)
        self.label_80 = QtWidgets.QLabel(self.tab_3)
        self.label_80.setGeometry(QtCore.QRect(620, 250, 71, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_80.setFont(font)
        self.label_80.setObjectName("label_80")
        self.label_81 = QtWidgets.QLabel(self.tab_3)
        self.label_81.setGeometry(QtCore.QRect(620, 10, 71, 16))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setBold(True)
        font.setWeight(75)
        self.label_81.setFont(font)
        self.label_81.setObjectName("label_81")
        self.IDPatientinInspect = QtWidgets.QLineEdit(self.tab_3)
        self.IDPatientinInspect.setGeometry(QtCore.QRect(230, 50, 113, 22))
        self.IDPatientinInspect.setObjectName("IDPatientinInspect")
        self.IDDoctorInInspect = QtWidgets.QLineEdit(self.tab_3)
        self.IDDoctorInInspect.setGeometry(QtCore.QRect(50, 50, 113, 22))
        self.IDDoctorInInspect.setObjectName("IDDoctorInInspect")
        self.dateTimeInInspect = QtWidgets.QDateEdit(self.tab_3)
        self.dateTimeInInspect.setGeometry(QtCore.QRect(80, 310, 110, 22))
        self.dateTimeInInspect.setObjectName("dateTimeInInspect")
        self.historyCost_3 = QtWidgets.QLabel(self.tab_3)
        self.historyCost_3.setGeometry(QtCore.QRect(20, 310, 51, 16))
        self.historyCost_3.setObjectName("historyCost_3")
        self.tabWidget.addTab(self.tab_3, "")
        self.tab_4 = QtWidgets.QWidget()
        self.tab_4.setObjectName("tab_4")
        self.checkBox_doctor = QtWidgets.QCheckBox(self.tab_4)
        self.checkBox_doctor.setGeometry(QtCore.QRect(330, 40, 81, 20))
        self.checkBox_doctor.setObjectName("checkBox_doctor")
        self.checkBox_Patient = QtWidgets.QCheckBox(self.tab_4)
        self.checkBox_Patient.setGeometry(QtCore.QRect(460, 40, 81, 20))
        self.checkBox_Patient.setObjectName("checkBox_Patient")
        self.label_85 = QtWidgets.QLabel(self.tab_4)
        self.label_85.setGeometry(QtCore.QRect(380, 10, 251, 16))
        self.label_85.setObjectName("label_85")
        self.exportDocButton = QtWidgets.QPushButton(self.tab_4)
        self.exportDocButton.setGeometry(QtCore.QRect(290, 70, 271, 51))
        self.exportDocButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.exportDocButton.setObjectName("exportDocButton")
        self.exportXLSButton = QtWidgets.QPushButton(self.tab_4)
        self.exportXLSButton.setGeometry(QtCore.QRect(290, 130, 271, 51))
        self.exportXLSButton.setStyleSheet("QPushButton { background: rgb(170, 170, 255)}")
        self.exportXLSButton.setObjectName("exportXLSButton")
        self.tabWidget.addTab(self.tab_4, "")

        self.retranslateUi(Dialog)
        self.tabWidget.setCurrentIndex(2)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.label.setText(_translate("Dialog", "Add Doctor"))
        self.label_2.setText(_translate("Dialog", "Doctor Name"))
        self.label_3.setText(_translate("Dialog", "Specialty"))
        self.label_4.setText(_translate("Dialog", "Cost"))
        self.label_5.setText(_translate("Dialog", "Reception"))
        self.addUserButton.setText(_translate("Dialog", "Add"))
        self.updateUserButton.setText(_translate("Dialog", "Update"))
        self.label_6.setText(_translate("Dialog", "Update Doctor"))
        self.deleteUserButton.setText(_translate("Dialog", "Delete"))
        self.label_17.setText(_translate("Dialog", "Salaries"))
        self.label_8.setText(_translate("Dialog", "Reception"))
        self.label_18.setText(_translate("Dialog", "Salaries"))
        self.label_9.setText(_translate("Dialog", "Cost"))
        self.label_10.setText(_translate("Dialog", "Doctor Name"))
        self.label_11.setText(_translate("Dialog", "Specialty"))
        self.label_12.setText(_translate("Dialog", "ID"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab), _translate("Dialog", "Doctor"))
        self.label_50.setText(_translate("Dialog", "Name"))
        self.label_52.setText(_translate("Dialog", "Adres"))
        self.deletePatientButton.setText(_translate("Dialog", "Delete"))
        self.label_51.setText(_translate("Dialog", "Departure"))
        self.updatePatientButton.setText(_translate("Dialog", "Update"))
        self.label_49.setText(_translate("Dialog", "Add Patients"))
        self.addPatientButton.setText(_translate("Dialog", "Add"))
        self.label_53.setText(_translate("Dialog", "Birthday"))
        self.label_54.setText(_translate("Dialog", "Name"))
        self.label_56.setText(_translate("Dialog", "Update Patients"))
        self.label_57.setText(_translate("Dialog", "Adres"))
        self.label_58.setText(_translate("Dialog", "Birthday"))
        self.label_59.setText(_translate("Dialog", "Departure"))
        self.label_55.setText(_translate("Dialog", "ID"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_2), _translate("Dialog", "Patient"))
        self.histiroyIdDoctor.setText(_translate("Dialog", "ID"))
        self.historyNamePatient.setText(_translate("Dialog", ""))
        self.historySpecialtyDoctor.setText(_translate("Dialog", ""))
        self.histiroyNameDoctor.setText(_translate("Dialog", ""))
        self.historyCost.setText(_translate("Dialog", ""))
        self.historyPatientBirthday.setText(_translate("Dialog", ""))
        self.historyCost_2.setText(_translate("Dialog", ""))
        self.historyPatientAdress.setText(_translate("Dialog", ""))
        self.historyPatientSurname.setText(_translate("Dialog", ""))
        self.selectDatas.setText(_translate("Dialog", "Select"))
        self.label_78.setText(_translate("Dialog", "Doctor"))
        self.historyReception.setText(_translate("Dialog", ""))
        self.label_79.setText(_translate("Dialog", "Patient"))
        self.historyIdPatient.setText(_translate("Dialog", "ID"))
        self.saveData.setText(_translate("Dialog", "Save"))
        self.label_80.setText(_translate("Dialog", "Patient"))
        self.label_81.setText(_translate("Dialog", "Doctor"))
        self.historyCost_3.setText(_translate("Dialog", "Date:"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_3), _translate("Dialog", "Inspect"))
        self.checkBox_doctor.setText(_translate("Dialog", "Doctor"))
        self.checkBox_Patient.setText(_translate("Dialog", "Patient"))
        self.label_85.setText(_translate("Dialog", "Select tables to export"))
        self.exportDocButton.setText(_translate("Dialog", "Export as .docx"))
        self.exportXLSButton.setText(_translate("Dialog", "Export as .xlsx"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_4), _translate("Dialog", "Export"))

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Dialog = QtWidgets.QDialog()
    ui = Ui_Dialog()
    ui.setupUi(Dialog)
    Dialog.show()
    
    def tableDoctors():
        ui.tableDoctor.setRowCount(0)
        try:
            data = controller.Doctor.getDoctor()

            ui.tableDoctor.setColumnCount(6)
            header_labels = ['ID', 'Doctor Name', 'Doctor Specialty', 'Cost', 'Reception', 'Salaries']
            ui.tableDoctor.setHorizontalHeaderLabels(header_labels)
            for row in data:
                inx = data.index(row)
                ui.tableDoctor.insertRow(inx)
                ui.tableDoctor.setItem(inx,0, QtWidgets.QTableWidgetItem(str(row[0])))  
                ui.tableDoctor.setItem(inx,1, QtWidgets.QTableWidgetItem(str(row[1]))) 
                ui.tableDoctor.setItem(inx,2, QtWidgets.QTableWidgetItem(str(row[2]))) 
                ui.tableDoctor.setItem(inx,3, QtWidgets.QTableWidgetItem(str(row[3]))) 
                ui.tableDoctor.setItem(inx,4, QtWidgets.QTableWidgetItem(str(row[4])))
                ui.tableDoctor.setItem(inx,5, QtWidgets.QTableWidgetItem(str(row[5])))
                ui.tableDoctor.horizontalHeader().setStretchLastSection(True) 
                ui.tableDoctor.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch) 

        except Exception as e:
            pass
    def tablePatients():
        ui.tablePatients.setRowCount(0)
        try:
            data = controller.Patients.GetPatients()

            ui.tablePatients.setColumnCount(6)
            header_labels = ['ID', 'Patient Name', 'Patient Surname', 'Birthday', 'Adress',""]
            ui.tablePatients.setHorizontalHeaderLabels(header_labels)
            for row in data:
                inx = data.index(row)
                ui.tablePatients.insertRow(inx)
                ui.tablePatients.setItem(inx,0, QtWidgets.QTableWidgetItem(str(row[0])))  
                ui.tablePatients.setItem(inx,1, QtWidgets.QTableWidgetItem(str(row[1]))) 
                ui.tablePatients.setItem(inx,2, QtWidgets.QTableWidgetItem(str(row[2]))) 
                ui.tablePatients.setItem(inx,3, QtWidgets.QTableWidgetItem(str(row[3]))) 
                ui.tablePatients.setItem(inx,4, QtWidgets.QTableWidgetItem(str(row[4])))
                ui.tablePatients.horizontalHeader().setStretchLastSection(True) 
                ui.tablePatients.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch) 

        except Exception as e:
            pass

    def Inspection():
        
        ui.tableWidget_4.setRowCount(0)
        try:
            data = controller.Doctor.getDoctor()

            ui.tableWidget_4.setColumnCount(6)
            header_labels = ['ID', 'Doctor Name', 'Doctor Specialty', 'Cost', 'Reception', 'Salaries']
            ui.tableWidget_4.setHorizontalHeaderLabels(header_labels)
            for row in data:
                inx = data.index(row)
                ui.tableWidget_4.insertRow(inx)
                ui.tableWidget_4.setItem(inx,0, QtWidgets.QTableWidgetItem(str(row[0])))  
                ui.tableWidget_4.setItem(inx,1, QtWidgets.QTableWidgetItem(str(row[1]))) 
                ui.tableWidget_4.setItem(inx,2, QtWidgets.QTableWidgetItem(str(row[2]))) 
                ui.tableWidget_4.setItem(inx,3, QtWidgets.QTableWidgetItem(str(row[3]))) 
                ui.tableWidget_4.setItem(inx,4, QtWidgets.QTableWidgetItem(str(row[4])))
                ui.tableWidget_4.setItem(inx,5, QtWidgets.QTableWidgetItem(str(row[5])))
                ui.tableWidget_4.horizontalHeader().setStretchLastSection(True) 
                ui.tableWidget_4.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch) 

        except Exception as e:
            pass
   
        ui.tableWidget_3.setRowCount(0)
        try:
            data = controller.Patients.GetPatients()

            ui.tableWidget_3.setColumnCount(6)
            header_labels = ['ID', 'Patient Name', 'Patient Surname', 'Birthday', 'Adress',""]
            ui.tableWidget_3.setHorizontalHeaderLabels(header_labels)
            for row in data:
                inx = data.index(row)
                ui.tableWidget_3.insertRow(inx)
                ui.tableWidget_3.setItem(inx,0, QtWidgets.QTableWidgetItem(str(row[0])))  
                ui.tableWidget_3.setItem(inx,1, QtWidgets.QTableWidgetItem(str(row[1]))) 
                ui.tableWidget_3.setItem(inx,2, QtWidgets.QTableWidgetItem(str(row[2]))) 
                ui.tableWidget_3.setItem(inx,3, QtWidgets.QTableWidgetItem(str(row[3]))) 
                ui.tableWidget_3.setItem(inx,4, QtWidgets.QTableWidgetItem(str(row[4])))
                ui.tableWidget_3.horizontalHeader().setStretchLastSection(True) 
                ui.tableWidget_3.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch) 

        except Exception as e:
            pass
    #Doctors Process Start

    def saveDoctor():
        if controller.Doctor.saveDoctor( ui.addDoctorName.text(),ui.addDoctorSpecialty.text(),ui.addCost.text(),ui.addNameOfReception.text(),ui.addPercentage.text()) == True:
            ui.tableDoctor.clear()
            tableDoctors()
            ui.addDoctorName.clear()
            ui.addDoctorSpecialty.clear()
            ui.addCost.clear()
            ui.addNameOfReception.clear()
            ui.addPercentage.clear()
        else:
            messagebox.showinfo("ERROR","YOU MUST BE INSERT INFORMATION")
    
    def updateDoctor():        
       if controller.Doctor.updateDoctor(ui.IDDoctor.text(),ui.updateDoctorName.text(),ui.updateDoctorSpecialty.text(),ui.updateCost.text(),ui.updateNameOfReception.text(),ui.updatePercentage.text()) == True:
           ui.tableDoctor.clear()
           tableDoctors()
           ui.updateDoctorName.clear()
           ui.updateDoctorSpecialty.clear()
           ui.updateCost.clear()
           ui.updateNameOfReception.clear()
           ui.addPercentage.clear()
    def deleteDoctor():
        if controller.Doctor.deleteDoctors(ui.IDDoctor.text()) == True:
            ui.tableDoctor.clear()
            tableDoctors()
            ui.IDDoctor.clear()
    ui.addUserButton.clicked.connect(saveDoctor)
    ui.updateUserButton.clicked.connect(updateDoctor)
    ui.deleteUserButton.clicked.connect(deleteDoctor)
    #Doctors Proccess Ended
    
    #Patient Proccess Start
    def patientAdd():
        if controller.Patients.SavePatients( ui.addPatientsName.text(), ui.addPatientsSurname.text(),ui.addPatientsBirthday.text(),ui.addPatientsAdress.text()) == True:
            ui.tablePatients.clear()
            tablePatients()
            ui.addPatientsName.clear()
            ui.addPatientsSurname.clear()
            ui.addPatientsBirthday.clear()
            ui.addPatientsAdress.clear()
        else:
            messagebox.showinfo("ERROR","YOU MUST BE INSERT INFORMATION")
    def patientDelete():
        if controller.Patients.deletePatients(ui.IDPatient.text()) == True:
            ui.tableDoctor.clear()
            tablePatients()
            ui.IDPatient.clear()
    def patientUpdate():
        if controller.Patients.UpdatePatients(ui.IDPatient.text(), ui.updatePatientsName.text(), ui.updatePatientsSurname.text(),ui.updatePatientsBirthday.text(),ui.updatePatientsAdress.text()) == True:
            ui.tablePatients.clear()
            tablePatients()
            ui.updatePatientsName.clear()
            ui.updatePatientsSurname.clear()
            ui.updatePatientsBirthday.clear()
            ui.updatePatientsAdress.clear()
        else:
            messagebox.showinfo("ERROR","YOU MUST BE INSERT INFORMATION")
    ui.addPatientButton.clicked.connect(patientAdd)
    ui.updatePatientButton.clicked.connect(patientUpdate)
    ui.deletePatientButton.clicked.connect(patientDelete)
    #Patient Proccess ended

    #Inspect Proccess Start

    def SelectDatas():
        
        try:
            dataDoctors = controller.NewInspection.getDoctor(ui.IDDoctorInInspect.text())
            
                
            dataPatient = controller.NewInspection.getPatients(ui.IDPatientinInspect.text())
            for i in dataDoctors:
                ui.histiroyNameDoctor.setText(i[1])
                ui.historySpecialtyDoctor.setText(i[2])
                ui.historyCost.setText(str(i[3]))
                ui.historyReception.setText ( i[4])

                ui.historyCost_2.setText(str(controller.NewInspection.makePrice(i[3],i[5])))
            for i in dataPatient:
                ui.historyNamePatient.setText (i[1])
                ui.historyPatientSurname.setText(i[2])
          #      ui.historyPatientBirthday.setText (dataPatient[3])
                ui.historyPatientAdress.setText (i[4])
        except Exception as e:
            print(e)
        """except:
            if ui.IDDoctorInInspect.text != int:
                messagebox.showinfo("ERROR","YOU MUST BE INSERT DOCTOR ID")
            if ui.IDPatientinInspect.text != int:
                messagebox.showinfo("ERROR","YOU MUST BE INSERT PATIENT ID")"""
    def SaveDatas():
           
        if controller.NewInspection.saveInspection(ui.IDDoctorInInspect.text(),ui.IDPatientinInspect.text(),ui.historyCost_2.text(),ui.dateTimeInInspect.text()) == True:
            ui.IDDoctorInInspect.clear()
            ui.IDPatientinInspect.clear()
            ui.histiroyNameDoctor.clear()
            ui.historySpecialtyDoctor.clear()
            ui.historyCost.clear()
            ui.historyReception.clear()
            ui.historyNamePatient.clear()
            ui.historyPatientSurname.clear()
            ui.historyPatientBirthday.clear()
            ui.historyPatientAdress.clear()
            ui.historyCost_2.clear()
            ui.dateTimeInInspect.clear()
    ui.saveData.clicked.connect(SaveDatas)
    ui.selectDatas.clicked.connect(SelectDatas)
    #Inspect Proccess End


     #Export Proccess Start
    def exportDoc():
        from datetime import date
        today = date.today()
        date = today.strftime("%d/%m/%Y")
        document = Document()
        document.add_heading(f'Hospital Export Data: {date}', 0)

        if (ui.checkBox_doctor.isChecked()):
            clientRecords = controller.Doctor.getDoctor()
            print(clientRecords)
            document.add_heading('Doctor table', level=1)
            document.add_paragraph(f'We have {len(clientRecords)} doctors', style='Intense Quote')
            table = document.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Id'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Specialty'
            hdr_cells[3].text = 'Cost'
            hdr_cells[4].text = 'Reception'
            hdr_cells[5].text = 'Percentage'
            for entry in clientRecords:
                row_cells = table.add_row().cells
                row_cells[0].text = str(entry[0])
                row_cells[1].text = str(entry[1])
                row_cells[2].text = str(entry[2])
                row_cells[3].text = str(entry[3])
                row_cells[4].text = str(entry[4])
                row_cells[5].text = str(entry[5])
            document.add_page_break()

        if (ui.checkBox_Patient.isChecked()):
            routeRecords = controller.Patients.GetPatients()
            document.add_heading('Patient table', level=1)
            document.add_paragraph(f'We have {len(routeRecords)} Patients', style='Intense Quote')
            table = document.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Id'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Surname'
            hdr_cells[3].text = 'Birth'
            hdr_cells[4].text = 'Adress'
            for entry in routeRecords:
                row_cells = table.add_row().cells
                row_cells[0].text = str(entry[0])
                row_cells[1].text = str(entry[1])
                row_cells[2].text = str(entry[2])
                row_cells[3].text = str(entry[3])
                row_cells[4].text = str(entry[4])
                
            document.add_page_break()

        try:
            document.save('{0}-data.doc'.format(today))
            document.save('{0}-data.docx'.format(today))
            messagebox.showinfo("Success","Document saved")
        except Exception as e:
            messagebox.showerror("Failed",e)
    def exportExcel():
        today = date.today()
    
        if (ui.checkBox_doctor.isChecked()):
            data = pandas.DataFrame(controller.Doctor.getDoctor(),columns=['ID','Name','Specialty','Cost','Reception','Percentage'])
            dataToExcel = pandas.ExcelWriter("Data-Doctor {0}.xlsx".format(today), engine='xlsxwriter')
            data.to_excel(dataToExcel)
            dataToExcel.save()
    
        if (ui.checkBox_Patient.isChecked()):
            data = pandas.DataFrame(controller.Patients.GetPatients(),columns=['ID','Name','Surname','Birthday','Adress'])
            dataToExcel = pandas.ExcelWriter("Data-Patients {0}.xlsx".format(today), engine='xlsxwriter')
            data.to_excel(dataToExcel)
            dataToExcel.save()
    ui.exportDocButton.clicked.connect(exportDoc)
    ui.exportXLSButton.clicked.connect(exportExcel)




    #Export Proccess End

   
tableDoctors()
tablePatients()
Inspection()
sys.exit(app.exec_())

    
    